import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

SRC = r'C:\Users\yusun\Desktop\1-2.프로젝트기획서_작성본.docx'
DST = r'C:\Users\yusun\Desktop\1-2.프로젝트기획서_작성본.docx'

d = Document(SRC)

# 정상 표 폭 (다른 표들 평균): col0=1648, col1=6743, total=8391 (dxa)
TARGET_WIDTHS = [1648, 6743]

t = d.tables[1]
tbl = t._tbl

# 1) tblGrid 수정
tblGrid = tbl.find(qn('w:tblGrid'))
gridCols = tblGrid.findall(qn('w:gridCol'))
for gc, w in zip(gridCols, TARGET_WIDTHS):
    gc.set(qn('w:w'), str(w))

# 2) 각 셀의 w:tcW 수정
for row in t.rows:
    cells_seen = set()
    for ci, cell in enumerate(row.cells):
        # avoid double-processing merged cells
        if id(cell._tc) in cells_seen:
            continue
        cells_seen.add(id(cell._tc))
        tcPr = cell._tc.find(qn('w:tcPr'))
        if tcPr is None:
            continue
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None and ci < len(TARGET_WIDTHS):
            tcW.set(qn('w:w'), str(TARGET_WIDTHS[ci]))
            tcW.set(qn('w:type'), 'dxa')

# 3) 표 layout을 fixed로 고정 (overflow 방지) — auto면 내용에 따라 늘어남
tblPr = tbl.find(qn('w:tblPr'))
if tblPr is not None:
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        from docx.oxml import OxmlElement
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    # tblW (전체 표 너비)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblW.set(qn('w:w'), str(sum(TARGET_WIDTHS)))
        tblW.set(qn('w:type'), 'dxa')

d.save(DST)
print('Table 1 width fixed to 8391 dxa (5.83 in)')

# 검증
d2 = Document(DST)
t2 = d2.tables[1]
tblGrid2 = t2._tbl.find(qn('w:tblGrid'))
widths = [int(gc.get(qn('w:w'))) for gc in tblGrid2.findall(qn('w:gridCol'))]
print(f'verified widths: {widths}, total: {sum(widths)} dxa = {sum(widths) * 635 / 914400:.2f} in')
