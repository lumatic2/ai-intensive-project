import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\yusun\Desktop\1-2.프로젝트기획서_작성본.docx'
DST = SRC

d = Document(SRC)

def set_keep_next(p):
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    keepNext = pPr.find(qn('w:keepNext'))
    if keepNext is None:
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)

# 모든 body level 단락에 keep_next 적용 (빈 단락 포함)
# 단, 마지막 단락은 제외 (불필요)
body = d.element.body
all_paragraphs = list(d.paragraphs)
count_total = 0
count_empty = 0
for p in all_paragraphs[:-1]:  # 마지막 제외
    set_keep_next(p)
    count_total += 1
    if not p.text.strip():
        count_empty += 1

# 추가: 표의 첫 행에 cantSplit 적용 (행 자체가 페이지 경계에서 안 쪼개지게)
for t in d.tables:
    for row in t.rows:
        trPr = row._tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row._tr.insert(0, trPr)
        cantSplit = trPr.find(qn('w:cantSplit'))
        if cantSplit is None:
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)

d.save(DST)
print(f'keep_next: {count_total} paragraphs (그 중 빈 단락 {count_empty}개 포함)')
print(f'모든 표 행에 cantSplit 적용')
