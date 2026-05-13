import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\yusun\Desktop\1-2.프로젝트기획서_작성본.docx'
DST = SRC

d = Document(SRC)
FONT = '맑은 고딕'

def set_run_font(run):
    """run의 ascii/eastAsia/hAnsi/cs 폰트를 모두 맑은 고딕으로."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    # remove existing rFonts
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is not None:
        rPr.remove(rFonts)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'), FONT)
    rPr.insert(0, rFonts)

def set_paragraph_keep_with_next(p):
    """문단을 다음 문단/표와 분리되지 않게 (keep with next)."""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    keepNext = pPr.find(qn('w:keepNext'))
    if keepNext is None:
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)

# ── 1. 모든 run의 폰트를 맑은 고딕으로 ──
run_count = 0
# 본문 paragraphs
for p in d.paragraphs:
    for r in p.runs:
        set_run_font(r)
        run_count += 1
# 표 내부 paragraphs
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r)
                    run_count += 1
print(f'폰트 적용: {run_count} runs → 맑은 고딕')

# ── 2. (n) 제목 단락에 keep-with-next 적용 ──
# 본문 paragraphs 중 텍스트가 비어있지 않은 것들에 keep_next.
# 표 사이에 끼인 "(1) ...", "(2) ..." 같은 제목 단락이 항상 다음 표와 붙도록.
keep_count = 0
for p in d.paragraphs:
    if p.text.strip():
        set_paragraph_keep_with_next(p)
        keep_count += 1
print(f'keep-with-next 적용: {keep_count} 단락')

d.save(DST)
print(f'saved: {DST}')
