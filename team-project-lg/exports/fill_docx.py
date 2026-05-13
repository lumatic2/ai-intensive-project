import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from copy import deepcopy

SRC = r'C:\Users\yusun\Desktop\1-2.프로젝트기획서.docx'
DST = r'C:\Users\yusun\Desktop\1-2.프로젝트기획서_작성본.docx'

d = Document(SRC)

def set_cell(cell, text):
    """Replace cell text preserving first paragraph's run style."""
    # clear extra paragraphs
    for p in list(cell.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    # clear runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # split by newline into multiple runs with line breaks
    lines = text.split('\n')
    for i, line in enumerate(lines):
        r = p.add_run(line)
        if i < len(lines) - 1:
            r.add_break()

# ==================== Table 0: 브레인스토밍 ====================
T = d.tables[0]
# rows 0-2 already filled. row 3 header. rows 4+ = 후보 주제
brainstorm = [
    ("AI Second Brain OS (Obsidian 기반)", "전유성",
     "개인 지식 관리에 LLM을 결합한 시스템. LG 가전 멘토링 트랙과 연결성이 약함."),
    ("AI 페르소나 토론 시스템", "전유성",
     "다중 LLM 에이전트 활용 사례. 데모 임팩트가 가전 도메인보다 약함."),
    ("AI 의사결정 보조 시스템", "전유성",
     "범용성은 높으나 LG 멘토 산업 도메인과 거리감."),
    ("로봇청소기 Physical AI 시뮬레이터", "전유성",
     "LG 가전 멘토 연결성 강함. Physical AI·Explainable AI 트렌드와 결합. 3주 MVP 구현 가능 → 최종 채택."),
]
for i, (subj, who, why) in enumerate(brainstorm):
    r = 4 + i
    set_cell(T.rows[r].cells[0], subj)
    set_cell(T.rows[r].cells[2], who)
    set_cell(T.rows[r].cells[3], why)

# ==================== Table 1: 아이디어 기획서 ====================
T = d.tables[1]
# row 1: 팀원 및 역할 — already partially filled by user, append rest
set_cell(T.rows[1].cells[1],
    "▪팀장: 전유성 (총괄·기획·일정 관리, 백엔드·LLM 통합, 발표 스토리)\n"
    "▪팀원: 김준성 (시장 조사, 공개 IoT 데이터셋 분석, Mock dataset 설계, PPT)\n"
    "▪팀원: 박주상 (ML 이벤트 분류 모델 학습, LLM 프롬프트, 상황 추론 로직)")
# row 3: 제안 배경 및 필요성
set_cell(T.rows[3].cells[1],
    "▪제안 배경: 로봇청소기 보급률은 높아졌으나 대부분 \"정해진 시간 청소\" 수준에 머무름. "
    "LLM·Agentic AI 발전으로 가전이 \"명령 수행 기계\"에서 \"상황 이해 파트너\"로 진화 가능한 시점.\n"
    "▪필요성: 현재 로봇청소기는 공간은 인식하지만 상황·맥락을 이해하지 못함. "
    "사용자가 AI 가전의 자율 행동을 신뢰하려면 판단 이유가 설명되어야 함 (Explainable AI UX).")
# row 4: 유사 제품 현황 및 비교
set_cell(T.rows[4].cells[1],
    "▪LG CodeZero AI: SLAM·장애물 인식 강점. 상황 맥락 추론·자연어 설명 없음.\n"
    "▪삼성 비스포크 제트봇: 객체 인식 카메라. 이벤트 기반 우선순위 변경 없음.\n"
    "▪로보락 S8 Pro Ultra: 자동 비움·물걸레. 사용자가 명령 입력 필요.\n"
    "▪에코백스 X2 Omni: 듀얼 카메라. 시간 예약 중심.")
# row 5: 제안 내용
set_cell(T.rows[5].cells[1],
    "▪개발 목표: 생활 이벤트(귀가·요리·취침·손님 방문)와 외부 상황(날씨·시간)을 입력받아 "
    "공간별 청소 우선순위를 산출하고, AI가 그 이유를 자연어로 설명하는 시뮬레이터 구현.\n"
    "▪개발 내용:\n"
    "1. 2D 집 맵 UI + 공간별 속성(오염도·소음 민감도·최근 청소 시각) 정의\n"
    "2. 시나리오 이벤트 입력 → Rule-based scoring으로 공간별 priority score 계산\n"
    "3. LLM이 현재 상황을 요약하고 청소 전략을 자연어로 설명\n"
    "4. 공개 IoT 데이터셋 분석으로 가중치 보정 + scikit-learn ML 이벤트 분류기 학습")
# row 6: 수행 방법
set_cell(T.rows[6].cells[1],
    "▪데이터 확보 방안: 공개 IoT/스마트홈 데이터셋(UCI Activities of Daily Living, Kaggle Smart Home) "
    "분석으로 시간대별 공간 사용 패턴 추출. Mock dataset은 분석 결과를 반영해 직접 설계.\n"
    "▪추진 전략: 1주차 설계·데이터셋 분석 / 2주차 백엔드·LLM·ML 모델 구현 / 3주차 시나리오 통합·발표 준비. "
    "Python FastAPI 백엔드 + Streamlit 프론트엔드 + OpenAI/Claude API.")
# row 7: 기대효과 및 활용방안
set_cell(T.rows[7].cells[1],
    "▪기대효과: LG 가전의 차세대 AI UX 방향성 제시. Explainable AI를 가전 도메인에 적용한 PoC 확보. "
    "Home Agent / Home OS로 확장 가능한 기초 아키텍처 구축.\n"
    "▪활용방안: LG 로봇청소기 차세대 펌웨어 UX 기획 자료, ThinQ 앱 \"AI 청소 일지\" 기능 통합, "
    "에어컨·공기청정기·냉장고 등 다른 가전으로 동일 의사결정 패턴 확장.")

# ==================== Table 2: PEST/STEEP ====================
T = d.tables[2]
pest = {
    "정책적 배경": "▪정부 AI 산업 육성 정책 강화 (K-AI 전략, AI 기본법 시행). "
                "개인정보보호법 강화로 On-device AI·설명 가능한 AI(XAI) 수요 증가.",
    "경제적 배경": "▪글로벌 로봇청소기 시장 2024년 약 80억$, 2030년 200억$ 전망 (CAGR 14%). "
                "국내 점유율 경쟁 격화 (LG·삼성·로보락·에코백스). AI 가전 프리미엄이 신성장 동력.",
    "사회적 배경": "▪1인 가구·맞벌이·반려동물 가구 증가로 자율형 가전 수요 증가. "
                "사용자가 AI 자율 행동을 신뢰하지 못해 \"수동 모드\"로 회귀하는 경향 보고됨.",
    "트렌드 배경": "▪Physical AI (NVIDIA Project GR00T, Tesla Optimus, Figure 02). "
                "Agentic AI (사용자 명령 없이 스스로 계획·실행). Smart Home OS 흐름 정점.",
    "기술적 배경": "▪LLM 추론·설명 능력이 실시간 응답 가능 수준 (GPT-4o, Claude Opus, Gemini). "
                "클라우드 API로 가전+LLM 결합이 현실적 비용. Rule-based + LLM 하이브리드 아키텍처 검증됨.",
    "제도적 배경": "▪AI 기본법·개인정보보호법 강화 → 가전 AI 의사결정의 설명 가능성·투명성 요구 증가. "
                "Explainable AI가 제도적 요구로 자리잡는 추세.",
}
for r, row in enumerate(T.rows):
    key = row.cells[0].text.strip()
    if key in pest:
        set_cell(row.cells[1], pest[key])

# ==================== Table 3: 3C ====================
T = d.tables[3]
c3 = {
    "제안자 능력": "▪개발 능력: Python·FastAPI·React/Streamlit·LLM API 활용 (팀장 전유성).\n"
              "▪AI 모델 이해: 머신러닝·딥러닝·프롬프트 엔지니어링 (박주상, 인공지능 전공).\n"
              "▪시장·사용자 분석: 시장 조사·UX 설계 (김준성, 글로벌경영 전공).\n"
              "▪선행 경험: AI Intensive 1주차에서 OpenAI API·Whisper·TTS 활용 음성 비서 구현 완료.",
    "경쟁제품/기술/특허 분석": "▪LG CodeZero AI: SLAM·장애물 인식 강점. 상황 맥락 추론·자연어 설명 없음.\n"
                       "▪삼성 비스포크 제트봇: 객체 인식 카메라, SmartThings 통합. 이벤트 기반 우선순위 변경 없음.\n"
                       "▪로보락 S8 Pro Ultra: 자동 비움·물걸레. 사용자가 명령 입력 필요.\n"
                       "▪에코백스 X2 Omni: 듀얼 카메라. 시간 예약 중심.\n"
                       "▪공통점: \"왜 그 행동을 했는지\"를 설명하는 가전은 없음.",
    "고객 분석": "▪1차 타겟: 30~40대 맞벌이·1인 가구 — 로봇청소기 사용 경험은 있으나 \"예측 불가능성\"에 불만.\n"
            "▪2차 타겟: 반려동물·영유아 가구 — 청소 타이밍이 생활 리듬에 민감.\n"
            "▪공통 페인포인트: \"왜 지금 청소하지?\", \"왜 저 방을 안 청소했지?\" — AI 의사결정의 불투명성.",
}
for row in T.rows:
    key = row.cells[0].text.strip()
    if key in c3:
        set_cell(row.cells[1], c3[key])

# ==================== Table 4: 필요성·차별성·준비사항·기대효과·활용방안 (각 한 줄) ====================
T = d.tables[4]
m4 = {
    "필요성": "▪현재 로봇청소기는 공간은 알지만 상황·맥락을 이해하지 못해 사용자 신뢰를 얻지 못하는 한계가 있으므로, "
          "AI 기반 상황 추론과 설명 가능한 의사결정 UX가 필요하다.",
    "차별성": "▪청소 경로·성능 최적화가 아닌 AI 가전의 의사결정 UX에 집중하며, Rule-based scoring + LLM 자연어 설명을 "
          "결합한 하이브리드 아키텍처로 \"왜 그렇게 청소했는지\"를 사용자에게 설명할 수 있다.",
    "현재까지": "▪5-Layer 시스템 아키텍처 설계, 4개 시나리오(비 오는 날 귀가·요리 직후·취침 직전·손님 방문 전) 정의, "
            "Mock dataset 컬럼 구조와 scoring 가중치 초안, OpenAI API 통합 경험 확보 완료.",
    "기대효과": "▪LG 가전의 차세대 AI UX 방향성을 제시하고, Explainable AI를 가전 도메인에 적용한 PoC를 확보하여 "
           "Home Agent / Home OS로 확장 가능한 기초 아키텍처를 구축한다.",
    "활용방안": "▪LG 로봇청소기 차세대 펌웨어 UX 기획 자료 및 ThinQ 앱 \"AI 청소 일지\" 기능으로 통합 가능하며, "
           "에어컨·공기청정기·냉장고 등 다른 가전으로 동일한 의사결정 패턴을 확장 적용할 수 있다.",
}
for row in T.rows:
    key = row.cells[0].text.strip().replace(' ', '').replace('/', '')
    for k, v in m4.items():
        if k.replace(' ', '') in key:
            set_cell(row.cells[1], v)
            break

# ==================== Table 5: 개발 목표·내용 ====================
T = d.tables[5]
set_cell(T.rows[0].cells[1],
    "▪생활 이벤트(귀가·요리·취침·손님 방문)와 외부 상황(날씨·시간대)을 입력받아 공간별 청소 우선순위를 산출.\n"
    "▪AI가 의사결정 이유를 자연어로 설명.\n"
    "▪4개 이상의 시나리오에서 우선순위가 일관되게 변화함을 작동 데모로 입증.")
set_cell(T.rows[1].cells[1],
    "▪Spatial Layer: 집 구조 2D 맵 + 공간별 속성(오염도·사용 빈도·청소 민감도·소음 민감도·최근 청소 시각).\n"
    "▪Behavioral Layer: 이벤트 정의(귀가·요리·취침·운동 후·손님 방문) + 공간별 가중치 매핑.\n"
    "▪Context Layer: 시간·이벤트·공간 상태·날씨·청소 기록을 LLM이 자연어 컨텍스트로 재구성.\n"
    "▪Decision Layer: Rule-based scoring으로 공간별 priority score 계산 → 청소 순서·제외·모드 결정.\n"
    "▪Explainable AI Layer: LLM이 의사결정 이유를 자연어로 생성 (\"왜 먼저?\", \"왜 제외?\").\n"
    "▪데이터 기반 보강: 공개 IoT 데이터셋(UCI ADL, Kaggle Smart Home) 분석으로 가중치 보정 + scikit-learn ML 이벤트 분류기 학습.")

# ==================== Table 6: 달성 목표·전략 ====================
T = d.tables[6]
set_cell(T.rows[0].cells[1],
    "▪4개 시나리오(비 오는 날 귀가·요리 직후·취침 직전·손님 방문 전)에서 공간별 청소 우선순위가 "
    "상황에 따라 다르게 산출되고, AI가 그 이유를 자연어로 설명하는 작동 데모 완성.\n"
    "▪발표 시 멘토·심사자가 \"이것이 LG 가전의 미래 UX가 될 수 있다\"고 납득할 서사 구성.")
set_cell(T.rows[1].cells[1],
    "1. 1주차 (5/13~5/19) — 설계·기초: Mock dataset·시나리오·scoring rule 명세 / 백엔드·프론트 스켈레톤 / "
    "공개 IoT 데이터셋 1개 선정·기초 분석.\n"
    "2. 2주차 (5/20~5/26) — 핵심 구현: Scoring engine / LLM 프롬프트·API 연결 / 2D 맵 UI / "
    "ML 이벤트 분류 모델 학습·평가 / 시나리오 1~2 작동 검증.\n"
    "3. 3주차 (5/27~5/30) — 통합·발표: 시나리오 3~4 통합 / AI 설명 UX 다듬기 / 발표 PPT·데모 영상·시연 리허설.")

# ==================== Table 7: 개발 일정 (7행 13열, 1~12 컬럼) ====================
# 추진내용 6개 행이 미리 정해져 있음. 우리 일정에 맞게 텍스트 교체 + ● 표시
T = d.tables[7]
schedule_rows = [
    ("프로젝트 계획 및 보고",            [1, 2, 7, 12]),
    ("기존 제품 정밀 분석 · 데이터셋 분석",  [1, 2, 3, 4]),
    ("시스템 설계 · scoring rule 명세",   [2, 3, 4]),
    ("백엔드 · LLM · ML 모델 구현",      [4, 5, 6, 7, 8]),
    ("프론트엔드 · 시나리오 통합",        [7, 8, 9, 10]),
    ("발표 PPT · 데모 영상 · 리허설",    [10, 11, 12]),
]
for i, (label, marks) in enumerate(schedule_rows):
    r = 1 + i
    if r < len(T.rows):
        set_cell(T.rows[r].cells[0], label)
        for c in range(1, 13):
            set_cell(T.rows[r].cells[c], "●" if c in marks else "")

# ==================== Table 8: 참여 인원 ====================
T = d.tables[8]
members = [
    ("전유성 (팀장)", "▪프로젝트 총괄·기획·일정 관리 / 백엔드(FastAPI)·LLM 통합 / 발표 스토리·멘토 커뮤니케이션. "
                  "AI Intensive 1주차 OpenAI API·Whisper·TTS 음성 비서 구현 완료."),
    ("김준성", "▪시장·경쟁 제품 조사 (LG·삼성·로보락·에코백스) / 공개 IoT 데이터셋 분석·가중치 보정 / "
            "Mock dataset 설계 / PPT 제작·시장성 파트 발표."),
    ("박주상", "▪ML 이벤트 분류 모델(scikit-learn) 학습·평가 / LLM 프롬프트 엔지니어링 / "
            "상황 추론·설명 생성 로직 / Scoring engine 보조 개발."),
]
for i, (name, role) in enumerate(members):
    r = 1 + i
    if r < len(T.rows):
        set_cell(T.rows[r].cells[0], name)
        set_cell(T.rows[r].cells[1], role)

# ==================== Table 9: 시장 분석 ====================
T = d.tables[9]
# row 1: 국내시장 규모 및 현황
set_cell(T.rows[1].cells[1],
    "▪글로벌 로봇청소기 시장: 2024년 약 80억$, 2030년 약 200억$ 전망 (CAGR 14%).\n"
    "▪국내 주요 제품: LG CodeZero Objet Master (AI 객체 인식·습식 청소), "
    "삼성 비스포크 제트봇 AI (카메라 기반·SmartThings 통합), 로보락 (가성비·자동 비움), 에코백스 X2 Omni (듀얼 카메라).\n"
    "▪시장 기회: 모든 경쟁사가 청소 성능·하드웨어로 경쟁 중. "
    "AI 의사결정의 설명 가능성·상황 이해는 미개척 영역으로 LG가 선점할 수 있는 차세대 UX 표준 영역.")

# ==================== Table 10: SWOT ====================
T = d.tables[10]
# layout: [0,0]=S내용 [0,1]=S [0,2]=W [0,3]=W내용 / [1,0]=O내용 [1,1]=O [1,2]=T [1,3]=T내용
set_cell(T.rows[0].cells[0],
    "▪Rule-based + LLM 하이브리드 아키텍처로 의사결정의 설명 가능성 확보.\n"
    "▪Explainable AI를 가전 도메인에 적용한 차별화된 접근.\n"
    "▪Mock data 기반 시뮬레이션으로 실 데이터 부재를 우회하면서도 작동 데모 구현.")
set_cell(T.rows[0].cells[3],
    "▪실 로봇청소기 데이터 접근 불가 → Mock data 의존.\n"
    "▪3주의 짧은 개발 기간.\n"
    "▪실제 SLAM·하드웨어 제어는 범위 외.")
set_cell(T.rows[1].cells[0],
    "▪Physical AI·Agentic AI·Explainable AI 트렌드 정점.\n"
    "▪LG·삼성 모두 차세대 AI 가전 UX 모색 중.\n"
    "▪멘토 도메인과 직접 연결되어 발표 임팩트 확보.")
set_cell(T.rows[1].cells[3],
    "▪\"컨셉 PPT\"로 끝나면 차별성 약화 → 작동 데모 필수.\n"
    "▪범위 과확장 시 3주 안에 미완성.\n"
    "▪경쟁 팀 중 유사 주제 출현 가능성.")

# ==================== Table 11: STP ====================
T = d.tables[11]
stp = {
    "고객 분류": "▪A: 정해진 시간 자동 청소형 (기존 사용자 다수)\n"
            "▪B: 상황별 수동 호출형 (불만족 사용자)\n"
            "▪C: AI 자율 청소를 원하지만 신뢰 못 함 (잠재 타겟)",
    "목표 고객": "▪1차 타겟: C 세그먼트 — AI 자율성을 원하지만 예측 불가능성에 불만인 사용자.\n"
            "▪인구통계: 30~40대 맞벌이 / 1인 가구 / 반려동물·영유아 가구.\n"
            "▪특성: AI 가전에 호의적이나 \"왜?\"에 대한 답을 원함.",
    "시장에서": "▪\"청소를 잘하는 로봇\"이 아니라 \"왜 이렇게 청소했는지 설명하는 로봇\".\n"
            "▪\"단순 자동화\"에서 \"상황 이해·설명 가능한 Physical AI Agent\"로 시장 포지션 이동.",
}
for row in T.rows:
    key = row.cells[0].text.strip().replace(' ', '')
    for k, v in stp.items():
        if k.replace(' ', '') in key:
            set_cell(row.cells[1], v)
            break

# ==================== Table 12: 4P ====================
T = d.tables[12]
p4 = {
    "제품": "▪LG 로봇청소기 차세대 펌웨어 UX 컨셉 — 상황 기반 청소 + Explainable AI 설명 카드. "
        "ThinQ 앱 내 \"AI 청소 일지\" 기능으로 통합 가능. 향후 LG Home Agent OS의 핵심 인터페이스로 발전.",
    "가격": "▪펌웨어·앱 업데이트 형태로 기존 프리미엄 라인업에 무료 탑재. "
        "향후 구독형 \"AI Home Advisor\" 프리미엄 기능 (월 4,900원~9,900원) 확장 가능.",
    "판매 방법": "▪ThinQ 앱·LG 베스트샵·온라인 채널. "
            "B2B로 호텔·서비스드 레지던스 등 청소 자율성이 중요한 시장으로 확장.",
    "홍보 방법": "▪\"AI가 왜 그렇게 청소했는지 알려드립니다\" 캠페인. "
            "Physical AI·Explainable AI 트렌드 콘텐츠 마케팅. CES·IFA 등 가전 박람회 시연.",
}
for row in T.rows:
    key = row.cells[0].text.strip()
    for k, v in p4.items():
        if k in key:
            set_cell(row.cells[1], v)
            break

d.save(DST)
print(f"saved: {DST}")
