import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

SRC = r'C:\Users\yusun\Desktop\1-2.프로젝트기획서_작성본.docx'
DST = SRC

d = Document(SRC)

def set_cell(cell, text):
    for p in list(cell.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        r = p.add_run(line)
        # 폰트 맑은 고딕 유지
        rPr_run = r._element.find(qn('w:rPr'))
        if rPr_run is None:
            from docx.oxml import OxmlElement
            rPr_run = OxmlElement('w:rPr')
            r._element.insert(0, rPr_run)
        rFonts = rPr_run.find(qn('w:rFonts'))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts')
            rPr_run.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), '맑은 고딕')
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        rFonts.set(qn('w:hAnsi'), '맑은 고딕')
        rFonts.set(qn('w:cs'), '맑은 고딕')
        if i < len(lines) - 1:
            r.add_break()

# ───── 1. Table 1 (아이디어 기획서) 본문 압축 ─────
T = d.tables[1]

# row 3 제안 배경 및 필요성
set_cell(T.rows[3].cells[1],
    '▪제안 배경: 로봇청소기는 자동화되었으나 대부분 "정해진 시간 청소" 수준에 머무름. '
    'LLM·Agentic AI 발전으로 "명령 수행 기계"에서 "상황 이해 파트너"로 진화 가능한 시점.\n'
    '▪필요성: 사용자가 자율 행동을 신뢰하려면 AI 판단 이유가 설명되어야 함 (Explainable AI UX).')

# row 4 유사 제품 (이미 짧지만 정리)
set_cell(T.rows[4].cells[1],
    '▪LG CodeZero AI / 삼성 비스포크 제트봇 / 로보락 S8 / 에코백스 X2 — 모두 청소 성능·하드웨어로 경쟁 중.\n'
    '▪공통 한계: 상황 맥락 추론·자연어 설명·이벤트 기반 우선순위 변경 부재.')

# row 5 제안 내용
set_cell(T.rows[5].cells[1],
    '▪개발 목표: 생활 이벤트(귀가·요리·취침·손님 방문)와 외부 상황(날씨·시간)을 입력받아 '
    '공간별 청소 우선순위를 산출하고 AI가 그 이유를 자연어로 설명.\n'
    '▪개발 내용: 2D 집 맵 UI + Rule-based scoring + LLM 자연어 설명 + '
    '공개 IoT 데이터셋 분석으로 가중치 보정 + scikit-learn ML 이벤트 분류기 학습.')

# row 6 수행 방법
set_cell(T.rows[6].cells[1],
    '▪데이터: 공개 IoT 데이터셋(UCI ADL, Kaggle Smart Home) 분석 + Mock dataset 자체 설계.\n'
    '▪추진: 1주차 설계·데이터 분석 → 2주차 백엔드·LLM·ML 구현 → 3주차 시나리오 통합·발표. '
    'Python FastAPI + Streamlit + OpenAI/Claude API.')

# row 7 기대효과 및 활용방안
set_cell(T.rows[7].cells[1],
    '▪기대효과: LG 가전 차세대 AI UX 방향성 제시. Explainable AI 가전 PoC 확보. Home OS 확장 기반 아키텍처.\n'
    '▪활용방안: ThinQ 앱 통합 + 에어컨·공기청정기·냉장고 등 다른 가전에 동일 의사결정 패턴 확장.')

# ───── 2. Table 1 행 높이 최소값 축소 (atLeast 값을 작게 → 내용 따라 자연 축소) ─────
for ri, row in enumerate(T.rows):
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        continue
    trH = trPr.find(qn('w:trHeight'))
    if trH is not None:
        trH.set(qn('w:val'), '400')  # 최소 400 twips, 내용 따라 자동 확장

# ───── 3. 모든 헤딩의 pageBreakBefore 제거 (헤딩 이어서 흐르게) ─────
removed = 0
for p in d.paragraphs:
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        continue
    pbb = pPr.find(qn('w:pageBreakBefore'))
    if pbb is not None:
        pPr.remove(pbb)
        removed += 1

d.save(DST)
print(f'Table 1 본문 4개 셀 압축 완료')
print(f'Table 1 행 최소높이 400 twips로 축소')
print(f'헤딩 pageBreakBefore 제거: {removed}개')
