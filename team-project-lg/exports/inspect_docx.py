import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
d = Document(r'C:\Users\yusun\Desktop\1-2.프로젝트기획서.docx')
print('paragraphs:', len(d.paragraphs))
print('tables:', len(d.tables))
print()
# show first 30 paragraphs with index
for i, p in enumerate(d.paragraphs):
    print(f'P{i}: {p.text[:100]!r}')
print()
for ti, t in enumerate(d.tables):
    print(f'\n=== Table {ti}: {len(t.rows)}x{len(t.columns)} ===')
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip().replace('\n', ' / ')
            print(f'  [{ri},{ci}]: {txt[:80]!r}')
